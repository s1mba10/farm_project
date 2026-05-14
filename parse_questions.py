"""
Парсер вопросов из docx-теста MyTestXPro.

Особенности docx:
- В большинстве таблиц заголовочные строки имеют объединённые ячейки,
  и python-docx показывает len(cells)==3 с дублированным текстом.
- На разрывах страниц таблица разбивается на две: первая часть имеет
  одноячеечные строки (len(cells)==1), вторая часть начинается сразу
  с вариантов ответа.
- Поэтому мы накапливаем "текущий вопрос" и переключаемся на новый
  ТОЛЬКО когда встречаем "Задание №N".
"""
import json
import re
from docx import Document

SRC = "/mnt/user-data/uploads/ФАРМ_ВСЕ_ТЕМЫ.docx"
OUT = "/home/claude/questions.json"


def get_text(cell):
    return "\n".join(p.text for p in cell.paragraphs).strip()


def unique_cells(row):
    seen = []
    for c in row.cells:
        t = get_text(c)
        if not seen or seen[-1] != t:
            seen.append(t)
    return seen


def parse():
    doc = Document(SRC)
    questions = []
    current = None
    task_re = re.compile(r"Задание\s*№\s*(\d+)")

    for table in doc.tables:
        for row in table.rows:
            cells_unique = unique_cells(row)

            if len(cells_unique) == 1:
                txt = cells_unique[0]
                m = task_re.search(txt)
                if m:
                    if current:
                        questions.append(current)
                    current = {
                        "id": int(m.group(1)),
                        "question": "",
                        "type": "single",
                        "options": [],
                        "correct": [],
                    }
                    continue

                if current is None:
                    continue
                if not txt or txt == "---":
                    continue
                if "Выберите несколько" in txt:
                    current["type"] = "multi"
                    continue
                if re.match(r"^Выберите один из", txt):
                    current["type"] = "single"
                    continue
                if not current["question"]:
                    current["question"] = txt
                else:
                    if txt not in current["question"]:
                        current["question"] += " " + txt
                continue

            if len(cells_unique) >= 3 and current is not None:
                marker = cells_unique[1].strip()
                answer_text = cells_unique[2].strip()
                if marker in ("+", "-") and answer_text:
                    idx = len(current["options"])
                    current["options"].append(answer_text)
                    if marker == "+":
                        current["correct"].append(idx)

    if current:
        questions.append(current)

    questions = [q for q in questions if q["question"] and q["options"]]
    for q in questions:
        if len(q["correct"]) > 1:
            q["type"] = "multi"

    seen_ids = {}
    for q in questions:
        seen_ids[q["id"]] = q
    questions = sorted(seen_ids.values(), key=lambda x: x["id"])

    with open(OUT, "w", encoding="utf-8") as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)
    return questions


if __name__ == "__main__":
    qs = parse()
    print(f"Всего вопросов: {len(qs)}")
    print(f"  single: {sum(1 for q in qs if q['type']=='single')}")
    print(f"  multi:  {sum(1 for q in qs if q['type']=='multi')}")
    print(f"  без правильного ответа: {sum(1 for q in qs if not q['correct'])}")
    from collections import Counter
    opts_counter = Counter(len(q["options"]) for q in qs)
    print(f"  распределение по числу вариантов: {dict(opts_counter)}")
    ids = sorted(q["id"] for q in qs)
    missing = [i for i in range(1, max(ids) + 1) if i not in set(ids)]
    print(f"  макс id: {max(ids)}, пропуски: {len(missing)}")
    if missing[:10]:
        print(f"  первые пропуски: {missing[:10]}")
